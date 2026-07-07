import pytest
import tempfile
from pathlib import Path
from rag_src.doc_loader import UniversalDocLoader


def write_file(path: Path, content: str):
    path.write_text(content, encoding="utf-8")


def test_load_txt_file():
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        file = tmp_path / "sample.txt"
        write_file(file, "This is a plain text file.")

        loader = UniversalDocLoader(file)
        result = loader.load()

        assert "plain text" in result[0]


def test_load_md_file():
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        file = tmp_path / "README.md"
        write_file(file, "# Markdown Header\nSome content.")

        loader = UniversalDocLoader(file)
        result = loader.load()

        assert "Markdown Header" in result[0]
        assert "Some content" in result[0]


def test_load_html_file():
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        file = tmp_path / "page.html"
        write_file(file, "<html><body><h1>Hello</h1><p>World</p></body></html>")

        loader = UniversalDocLoader(file)
        result = loader.load()

        assert "Hello" in result[0]
        assert "World" in result[0]


def test_load_docx_file():
    import docx

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        file = tmp_path / "doc.docx"

        doc = docx.Document()
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")
        doc.save(str(file))

        loader = UniversalDocLoader(file)
        result = loader.load()

        assert "First paragraph" in result[0]
        assert "Second paragraph" in result[0]


def test_load_pdf_file():
    import fitz  # PyMuPDF

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        file = tmp_path / "doc.pdf"

        # Create a simple PDF
        pdf_doc = fitz.open()
        page = pdf_doc.new_page()
        page.insert_text((72, 72), "Hello PDF World!")
        pdf_doc.save(str(file))
        pdf_doc.close()

        loader = UniversalDocLoader(file)
        result = loader.load()

        assert "PDF World" in result[0]


def test_invalid_extension_raises():
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        file = tmp_path / "file.xyz"
        write_file(file, "unsupported content")

        with pytest.raises(ValueError, match="Unsupported file or directory"):
            UniversalDocLoader(file).load()
